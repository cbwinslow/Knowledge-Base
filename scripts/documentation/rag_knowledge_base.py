#!/usr/bin/env python3
"""
RAG Knowledge Base Manager
Integrates with existing knowledge base structure to provide semantic search and RAG capabilities
"""

import os
import sys
import json
import argparse
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import numpy as np

@dataclass
class Document:
    """Represents a document in the knowledge base"""
    file_path: str
    content: str
    metadata: Dict[str, Any]
    doc_id: str

class RAGKnowledgeBase:
    """RAG-enabled knowledge base manager"""
    
    def __init__(self, kb_path: str = "/home/cbwinslow/Knowledge-Base", 
                 persist_directory: str = "./chroma_db"):
        self.kb_path = Path(kb_path)
        self.persist_directory = persist_directory
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection = self.client.get_or_create_collection(
            name="knowledge_base",
            metadata={"hnsw:space": "cosine"}
        )
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text content from various file types"""
        try:
            if file_path.suffix.lower() in ['.md', '.txt', '.py', '.sh', '.yaml', '.yml', '.json']:
                return file_path.read_text(encoding='utf-8', errors='ignore')
            elif file_path.suffix.lower() == '.pdf':
                try:
                    import pypdf
                    with open(file_path, 'rb') as file:
                        reader = pypdf.PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    print("Warning: pypdf not installed, skipping PDF files")
                    return ""
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    return "\n".join([paragraph.text for paragraph in doc.paragraphs])
                except ImportError:
                    print("Warning: python-docx not installed, skipping Word documents")
                    return ""
            else:
                return ""
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""
    
    def get_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file path and content"""
        stat = file_path.stat()
        
        # Determine category based on path
        path_parts = file_path.relative_to(self.kb_path).parts
        category = path_parts[0] if len(path_parts) > 0 else "root"
        subcategory = path_parts[1] if len(path_parts) > 1 else ""
        
        # Determine file type
        file_type = file_path.suffix.lower().lstrip('.') or 'unknown'
        
        return {
            'file_path': str(file_path.relative_to(self.kb_path)),
            'absolute_path': str(file_path),
            'category': category,
            'subcategory': subcategory,
            'file_type': file_type,
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'indexed_at': datetime.now().isoformat()
        }
    
    def index_directory(self, directory: Optional[str] = None, 
                       file_patterns: List[str] = None) -> int:
        """Index documents from the knowledge base"""
        if file_patterns is None:
            file_patterns = ['*.md', '*.txt', '*.py', '*.sh', '*.yaml', '*.yml', 
                           '*.json', '*.pdf', '*.docx', '*.doc']
        
        if directory:
            search_path = self.kb_path / directory
        else:
            search_path = self.kb_path
        
        documents = []
        indexed_count = 0
        
        print(f"Indexing files in: {search_path}")
        
        for pattern in file_patterns:
            for file_path in search_path.rglob(pattern):
                if file_path.is_file():
                    # Skip certain directories
                    if any(skip in str(file_path) for skip in ['.git', '__pycache__', 'node_modules', '.venv']):
                        continue
                    
                    content = self.extract_text_from_file(file_path)
                    if content and len(content.strip()) > 50:  # Skip very short files
                        metadata = self.get_file_metadata(file_path)
                        doc_id = hashlib.md5(str(file_path).encode()).hexdigest()
                        
                        documents.append(Document(
                            file_path=str(file_path),
                            content=content,
                            metadata=metadata,
                            doc_id=doc_id
                        ))
        
        if documents:
            # Batch process embeddings
            print(f"Processing {len(documents)} documents...")
            contents = [doc.content for doc in documents]
            embeddings = self.embedding_model.encode(contents)
            
            # Add to ChromaDB
            self.collection.add(
                ids=[doc.doc_id for doc in documents],
                embeddings=embeddings.tolist(),
                documents=contents,
                metadatas=[doc.metadata for doc in documents]
            )
            
            indexed_count = len(documents)
        
        print(f"Indexed {indexed_count} new documents")
        return indexed_count
    
    def search(self, query: str, n_results: int = 5, 
               category_filter: Optional[str] = None) -> List[Dict[str, Any]]:
        """Search the knowledge base"""
        query_embedding = self.embedding_model.encode([query])
        
        where_filter = {"category": category_filter} if category_filter else None
        
        results = self.collection.query(
            query_embeddings=query_embedding.tolist(),
            n_results=n_results,
            where=where_filter,
            include=['documents', 'metadatas', 'distances']
        )
        
        search_results = []
        for i in range(len(results['ids'][0])):
            search_results.append({
                'doc_id': results['ids'][0][i],
                'content': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'similarity_score': 1 - results['distances'][0][i]  # Convert distance to similarity
            })
        
        return search_results
    
    def get_stats(self) -> Dict[str, Any]:
        """Get knowledge base statistics"""
        count = self.collection.count()
        
        # Get category distribution
        all_docs = self.collection.get(include=['metadatas'])
        categories = {}
        file_types = {}
        
        for metadata in all_docs['metadatas']:
            category = metadata.get('category', 'unknown')
            file_type = metadata.get('file_type', 'unknown')
            
            categories[category] = categories.get(category, 0) + 1
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': count,
            'categories': categories,
            'file_types': file_types,
            'persist_directory': self.persist_directory
        }
    
    def delete_by_category(self, category: str) -> int:
        """Delete all documents from a specific category"""
        self.collection.delete(where={"category": category})
        print(f"Deleted all documents from category: {category}")
        return 0
    
    def rebuild_index(self) -> int:
        """Rebuild the entire index"""
        # Clear existing collection
        self.client.delete_collection("knowledge_base")
        self.collection = self.client.get_or_create_collection(
            name="knowledge_base",
            metadata={"hnsw:space": "cosine"}
        )
        
        # Re-index everything
        return self.index_directory()

def main():
    parser = argparse.ArgumentParser(description='RAG Knowledge Base Manager')
    parser.add_argument('action', choices=['index', 'search', 'stats', 'rebuild', 'delete-category'],
                       help='Action to perform')
    parser.add_argument('--query', '-q', help='Search query')
    parser.add_argument('--directory', '-d', help='Directory to index (default: all)')
    parser.add_argument('--results', '-n', type=int, default=5, help='Number of search results')
    parser.add_argument('--category', '-c', help='Filter by category')
    parser.add_argument('--kb-path', default='/home/cbwinslow/Knowledge-Base', 
                       help='Knowledge base path')
    parser.add_argument('--persist-dir', default='./chroma_db', 
                       help='ChromaDB persist directory')
    
    args = parser.parse_args()
    
    rag_kb = RAGKnowledgeBase(args.kb_path, args.persist_dir)
    
    if args.action == 'index':
        count = rag_kb.index_directory(args.directory)
        print(f"Successfully indexed {count} documents")
    
    elif args.action == 'search':
        if not args.query:
            print("Error: --query is required for search")
            sys.exit(1)
        
        results = rag_kb.search(args.query, args.results, args.category)
        
        if not results:
            print("No results found")
        else:
            print(f"\nFound {len(results)} results for '{args.query}':\n")
            for i, result in enumerate(results, 1):
                print(f"{i}. {result['metadata']['file_path']}")
                print(f"   Category: {result['metadata']['category']}")
                print(f"   Similarity: {result['similarity_score']:.3f}")
                print(f"   Content: {result['content'][:200]}...")
                print()
    
    elif args.action == 'stats':
        stats = rag_kb.get_stats()
        print(f"\nKnowledge Base Statistics:")
        print(f"Total Documents: {stats['total_documents']}")
        print(f"\nCategories:")
        for category, count in sorted(stats['categories'].items()):
            print(f"  {category}: {count}")
        print(f"\nFile Types:")
        for file_type, count in sorted(stats['file_types'].items()):
            print(f"  {file_type}: {count}")
        print(f"\nPersist Directory: {stats['persist_directory']}")
    
    elif args.action == 'rebuild':
        count = rag_kb.rebuild_index()
        print(f"Rebuilt index with {count} documents")
    
    elif args.action == 'delete-category':
        if not args.category:
            print("Error: --category is required for delete-category")
            sys.exit(1)
        rag_kb.delete_by_category(args.category)

if __name__ == "__main__":
    main()